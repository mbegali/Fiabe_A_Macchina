import openai
import os
import pandas as pd
from docx import Document
from dotenv import load_dotenv

# Load API key from .env file
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Load the Excel file with grammar rules
df = pd.read_excel("explanation_database.xlsx")

# Function to search the Excel file for grammar rules
def search_grammar_rule(sentence):
    for _, row in df.iterrows():
        if row["Mistake"] in sentence:
            return f"Suggested correction: {row['Correction']} | Explanation: {row['Explanation']}"
    return None  # No match found

# Function to get feedback from GPT if the Excel file has no rule
def get_gpt_feedback(sentence):
    client = openai.OpenAI(api_key=OPENAI_API_KEY)
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are an expert Italian language teacher. Analyze the sentence for grammar mistakes and provide a correction with a short explanation in English."},
            {"role": "user", "content": f"Sentence: '{sentence}'"}
        ],
        temperature=0.2
    )
    return response.choices[0].message.content

# Load the student's Word document
doc = Document("student_essay.docx")
feedback_list = []  # Store all feedback for summary

# Process each paragraph
for para in doc.paragraphs:
    if para.text.strip():  # Ignore empty lines
        # Check the Excel file first
        grammar_feedback = search_grammar_rule(para.text)
        
        # If no rule found, use GPT
        if grammar_feedback is None:
            grammar_feedback = get_gpt_feedback(para.text)

        # Add feedback inline and save for summary
        feedback_list.append(f"- {para.text} ➝ {grammar_feedback}")
        para.text += f" [Attento qui: {grammar_feedback}]"

# Add a final summary at the end of the document
doc.add_paragraph("\n### Summary of Feedback ###")
for feedback in feedback_list:
    doc.add_paragraph(feedback)

# Save the document
doc.save("student_essay.docx")
print("Grammar feedback added with Excel + GPT!")
